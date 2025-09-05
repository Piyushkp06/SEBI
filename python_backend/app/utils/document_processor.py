import PyPDF2
from docx import Document
import os
import nltk
from nltk.tokenize import sent_tokenize
import spacy
from langdetect import detect
from typing import List, Dict, Any, Optional
import json
import logging
from pathlib import Path
import shutil
from .processor_config import ProcessorConfig

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def initialize_nlp():
    """Initialize spaCy with error handling and progress"""
    try:
        return spacy.load("en_core_web_sm")
    except OSError:
        logger.info("Downloading spaCy model... This may take a few minutes.")
        try:
            import subprocess
            import sys
            subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
            return spacy.load("en_core_web_sm")
        except Exception as e:
            logger.error(f"Failed to download spaCy model: {str(e)}")
            raise RuntimeError("Failed to initialize NLP model. Please run 'python -m spacy download en_core_web_sm' manually.")

# Initialize spaCy
nlp = initialize_nlp()

class DocumentProcessor:
    def __init__(self, config: ProcessorConfig | None = None):
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        self.config = config or ProcessorConfig()
        # Check for required dependencies
        self._check_dependencies()
        
    def _check_dependencies(self):
        """
        Check if all required external dependencies are installed
        """
        # No external dependencies required anymore
        pass
            
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF files using PyPDF2 only
        """
        try:
            # Check file size before processing
            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB limit
            if file_size > max_size:
                logger.error(f"PDF file too large ({file_size / 1024 / 1024:.1f}MB). Maximum size is 50MB")
                return ""

            # PDF text extraction
            text = ""
            with open(file_path, 'rb') as file:
                try:
                    reader = PyPDF2.PdfReader(file)
                    # Check number of pages
                    if len(reader.pages) > 100:
                        logger.warning(f"Large PDF detected ({len(reader.pages)} pages). Processing may take a while.")
                    
                    # Extract text page by page to manage memory
                    for page_num, page in enumerate(reader.pages):
                        try:
                            # Extract text from the page
                            page_text = page.extract_text() or ""
                            # Clean up the text and preserve some basic formatting
                            page_text = '\n'.join(line.strip() for line in page_text.splitlines() if line.strip())
                            text += page_text + "\n"
                            
                            # Free up memory
                            page_text = None
                            if (page_num + 1) % 10 == 0:
                                logger.info(f"Processed {page_num + 1} pages...")
                        except Exception as e:
                            logger.error(f"Error extracting text from page {page_num + 1}: {str(e)}")
                except Exception as e:
                    logger.error(f"Error reading PDF: {str(e)}")

            text = text.strip()
            if not text:
                logger.warning("No text could be extracted from the PDF. This might be a scanned or image-based PDF.")
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX files including tables and other elements
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT files with encoding detection and BOM handling
        """
        encodings = [
            'utf-8-sig',  # UTF-8 with BOM
            'utf-8',
            'latin-1',
            'ascii',
            'utf-16',
            'utf-16-le',
            'utf-16-be'
        ]
        
        # First try to detect the encoding
        try:
            import chardet
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected = chardet.detect(raw_data)
                if detected and detected.get('confidence', 0) > 0.7 and detected.get('encoding'):
                    detected_encoding = detected['encoding']
                    if isinstance(detected_encoding, str):
                        encodings.insert(0, detected_encoding)
        except ImportError:
            logger.warning("chardet not installed, falling back to default encoding list")
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    # Check if the content is readable
                    if not content.strip():
                        continue
                    logger.info(f"Successfully read file with {encoding} encoding")
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue
            
        logger.error(f"Could not decode text file with any encoding: {file_path}")
        return ""

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension
        """
        try:
            ext = Path(file_path).suffix.lower()
            if ext not in self.supported_extensions:
                logger.warning(f"Unsupported file extension: {ext}")
                return ""

            if ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                return self.extract_text_from_docx(file_path)
            elif ext == '.txt':
                return self.extract_text_from_txt(file_path)
            else:
                return ""
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return ""

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract structured information
        """
        text = self.extract_text(file_path)
        if not text:
            return {
                "success": False,
                "error": "No text could be extracted from the document"
            }

        try:
            # Detect language
            lang = detect(text)
            if lang != 'en':
                logger.warning(f"Document language detected as {lang}, not English")

            # Process with spaCy
            doc = nlp(text)

            # Extract key information
            extracted_info = {
                "success": True,
                "language": lang,
                "text": text,
                "sentences": sent_tokenize(text),
                "entities": self._extract_entities(doc),
                "investment_details": self._extract_investment_details(doc),
                "contact_info": self._extract_contact_info(doc),
                "key_phrases": self._extract_key_phrases(doc)
            }

            return extracted_info
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }

    def _extract_entities(self, doc) -> Dict[str, List[str]]:
        """
        Extract named entities from the document
        """
        entities = {
            "organizations": [],
            "people": [],
            "money": [],
            "percentages": [],
            "dates": []
        }

        for ent in doc.ents:
            if ent.label_ == "ORG":
                entities["organizations"].append(ent.text)
            elif ent.label_ == "PERSON":
                entities["people"].append(ent.text)
            elif ent.label_ == "MONEY":
                entities["money"].append(ent.text)
            elif ent.label_ == "PERCENT":
                entities["percentages"].append(ent.text)
            elif ent.label_ == "DATE":
                entities["dates"].append(ent.text)

        return entities

    def _extract_investment_details(self, doc) -> Dict[str, Any]:
        """
        Extract investment-related information
        """
        # Keywords for investment details
        return_keywords = ["return", "roi", "profit", "yield", "gain"]
        risk_keywords = ["risk", "guarantee", "assured", "guaranteed", "safe"]
        time_keywords = ["year", "month", "day", "term", "period"]

        investment_details = {
            "returns_mentioned": [],
            "risk_statements": [],
            "timeframes": [],
            "investment_amounts": []
        }

        for sent in doc.sents:
            sent_text = sent.text.lower()
            
            # Check for returns
            if any(keyword in sent_text for keyword in return_keywords):
                investment_details["returns_mentioned"].append(sent.text)

            # Check for risk statements
            if any(keyword in sent_text for keyword in risk_keywords):
                investment_details["risk_statements"].append(sent.text)

            # Check for timeframes
            if any(keyword in sent_text for keyword in time_keywords):
                investment_details["timeframes"].append(sent.text)

            # Extract monetary amounts
            for ent in sent.ents:
                if ent.label_ == "MONEY":
                    investment_details["investment_amounts"].append(ent.text)

        return investment_details

    def _extract_contact_info(self, doc) -> Dict[str, List[str]]:
        """
        Extract contact information from the document
        """
        import re
        
        contact_info = {
            "emails": [],
            "phones": [],
            "websites": []
        }

        text = doc.text

        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        contact_info["emails"] = re.findall(email_pattern, text)

        # Extract phone numbers
        phone_pattern = r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        contact_info["phones"] = re.findall(phone_pattern, text)

        # Extract websites
        website_pattern = r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b[-a-zA-Z0-9()@:%_\+.~#?&//=]*'
        contact_info["websites"] = re.findall(website_pattern, text)

        return contact_info

    def _extract_key_phrases(self, doc) -> List[str]:
        """
        Extract important phrases and statements
        """
        key_phrases = []
        
        # Important keywords to look for
        important_keywords = [
            "investment", "return", "guarantee", "profit", "opportunity",
            "limited time", "exclusive", "risk-free", "assured", "SEBI",
            "registered", "license", "registration"
        ]

        for sent in doc.sents:
            sent_text = sent.text.lower()
            if any(keyword in sent_text for keyword in important_keywords):
                key_phrases.append(sent.text.strip())

        return key_phrases
